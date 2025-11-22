# services/docx_service.py
from docx import Document
import os

OUTPUT_FOLDER = "output"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def create_docx(text: str, original_filename: str) -> str:
    """
    Create a DOCX from `text`. `original_filename` is used to derive the docx name.
    Returns the path to the created DOCX.
    """
    base = os.path.splitext(os.path.basename(original_filename))[0]
    out_path = os.path.join(OUTPUT_FOLDER, f"{base}.docx")
    doc = Document()
    for para in text.splitlines():
        if para.strip() == "":
            doc.add_paragraph("")  # preserve blank lines
        else:
            doc.add_paragraph(para)
    doc.save(out_path)
    return out_path

